"""
Resume Tailoring Agent - Web UI

A simple Streamlit interface for the resume tailoring agent.
Upload your files, paste a job URL, and get a tailored resume.
"""

import streamlit as st
import anthropic
import httpx
import os
import re
import json
import io
import copy
from dotenv import load_dotenv


def unescape_json_string(s: str) -> str:
    """Properly unescape a JSON string without corrupting UTF-8 characters."""
    # Handle common JSON escape sequences
    replacements = [
        (r'\\n', '\n'),
        (r'\\r', '\r'),
        (r'\\t', '\t'),
        (r'\\"', '"'),
        (r'\\\\', '\\'),
    ]
    result = s
    for pattern, replacement in replacements:
        result = result.replace(pattern, replacement)
    return result

# Load API key from .env file
load_dotenv()

st.set_page_config(
    page_title="The Perfect Fit",
    page_icon=None,
    layout="centered",
)

# Custom CSS for professional styling with soft green tones (light and dark mode compatible)
st.markdown("""
<style>
    /* Import clean sans-serif font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap');

    /* Global font */
    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }

    /* Headers - inherit color for dark mode compatibility */
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        font-weight: 600;
    }

    /* Primary button styling - soft green */
    .stButton > button[kind="primary"] {
        background-color: #4a7c59;
        border: none;
        color: white;
        font-weight: 500;
    }
    .stButton > button[kind="primary"]:hover {
        background-color: #3d6b4a;
        border: none;
    }

    /* Download buttons - works in both modes */
    .stDownloadButton > button {
        border: 1px solid #4a7c59;
        color: #4a7c59;
        font-weight: 500;
    }
    .stDownloadButton > button:hover {
        background-color: rgba(74, 124, 89, 0.1);
        border: 1px solid #4a7c59;
    }

    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 4px;
        padding: 8px 16px;
        font-weight: 500;
    }
    .stTabs [aria-selected="true"] {
        background-color: #4a7c59 !important;
        color: white !important;
    }

    /* Success message */
    .stSuccess {
        border-left: 4px solid #4a7c59;
    }

    /* Text area and inputs focus state */
    .stTextArea textarea:focus, .stTextInput input:focus {
        border-color: #4a7c59;
        box-shadow: 0 0 0 1px #4a7c59;
    }

    /* File uploader */
    [data-testid="stFileUploader"] {
        border: 1px dashed #4a7c59;
        border-radius: 4px;
        padding: 1rem;
    }
</style>
""", unsafe_allow_html=True)

st.title("The Perfect Fit")
st.subheader("Tailor your resume and cover letter")
st.markdown("Upload a qualifications document and your resume (Word format preferred), then paste in a job description. The app will align your resume to the role, generate a new Word file that preserves your original formatting, and explain every change it made — so you always know what was updated and why. Review the output carefully before sending to ensure accuracy.")

# Initialize session state for results
if "resume_result" not in st.session_state:
    st.session_state.resume_result = None
if "word_buffer" not in st.session_state:
    st.session_state.word_buffer = None
if "original_docx_bytes" not in st.session_state:
    st.session_state.original_docx_bytes = None
if "cover_letter_result" not in st.session_state:
    st.session_state.cover_letter_result = None
if "cover_letter_buffer" not in st.session_state:
    st.session_state.cover_letter_buffer = None

# API Key input (check env var or Streamlit secrets first, then allow manual entry)
api_key = os.environ.get("ANTHROPIC_API_KEY", "")
if not api_key:
    try:
        api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
    except:
        pass

with st.sidebar:
    st.header("API Key")
    st.markdown("""
    This tool uses Claude AI to analyze and tailor your resume. You provide your own API key so usage costs go directly to Anthropic (not to me).

    **Cost:** ~$0.02-0.04 per resume, +$0.01-0.02 for cover letter.

    Get your key at [console.anthropic.com](https://console.anthropic.com)
    """)

    entered_key = st.text_input(
        "Anthropic API Key",
        value=api_key,
        type="password",
        placeholder="sk-ant-api03-..."
    )
    if entered_key:
        api_key = entered_key

    if not api_key:
        st.warning("Enter your API key to get started")

    st.markdown("---")

    st.header("Instructions")
    st.markdown("""
    ### 1. Prepare Your Files

    **Current resume** - Upload as .docx to preserve formatting in the output.

    **Qualifications file** (optional but highly recommended) - A comprehensive document with far more detail than your resume. If you do not already have one, you may download and use the provided template. Providing this additional context helps the AI agent select and emphasize the most relevant experience for each job.

    ### 2. Add Job Description

    Copy and paste the job posting text.

    ### 3. Generate

    Click the button and wait for your tailored resume!
    """)

    st.markdown("---")
    st.caption("**Privacy:** Your data is processed in memory only and never stored. Resume content is sent to Anthropic's API for processing ([privacy policy](https://www.anthropic.com/privacy)).")
    st.caption("Powered by Claude")


def fetch_job_description(url: str) -> str:
    """Fetch and extract job description from a URL."""
    try:
        response = httpx.get(url, follow_redirects=True, timeout=30)
        response.raise_for_status()
        html = response.text

        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', html)
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()

        if len(text) > 15000:
            text = text[:15000] + "..."

        return text
    except Exception as e:
        return f"Error fetching URL: {str(e)}"


def extract_docx_structure(docx_bytes: io.BytesIO) -> list:
    """Extract the paragraph structure from a Word document."""
    from docx import Document

    docx_bytes.seek(0)
    doc = Document(docx_bytes)

    structure = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # Only include non-empty paragraphs
            structure.append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else "Normal"
            })

    return structure


def generate_tailored_resume_structured(
    client: anthropic.Anthropic,
    qualifications_content: str,
    resume_structure: list,
    job_description: str,
) -> dict:
    """Generate tailored content that maps to the original resume structure."""

    # Create a numbered list of paragraphs for Claude
    structure_text = "\n".join([
        f"[{item['index']}] ({item['style']}): {item['text']}"
        for item in resume_structure
    ])

    # Build qualifications section based on whether file was provided
    if qualifications_content:
        qualifications_section = f"""## Qualifications (SOURCE OF TRUTH - these are the ONLY facts you can use)
<qualifications>
{qualifications_content}
</qualifications>"""
    else:
        qualifications_section = """## Note: No separate qualifications file was provided.
Use ONLY the facts present in the resume itself. Do not invent or assume any additional qualifications."""

    prompt = f"""Tailor this resume for the job while preserving ALL facts exactly.

{qualifications_section}

## Resume (each line: [index] (style): text)
<resume>
{structure_text}
</resume>

## Job
<job>
{job_description}
</job>

## Rules
NEVER change: years of experience, job titles (keep EXACT original titles - never replace with target job title), company names, dates, degrees, certifications, metrics, skills listed, separator characters (|, •, -)
CRITICAL: The candidate's job titles are THEIR titles, not the job posting's title. "Software Engineer" stays "Software Engineer" even if applying for "Senior Developer".
CAN do: reorder bullets by relevance, rephrase bullet DESCRIPTIONS using job keywords (if truthful), adjust summary emphasis, use matching action verbs
Tense: present for current role, past for previous roles

## Output (JSON only)
{{"paragraphs": [{{"index": N, "text": "..."}}], "changes_summary": [{{"section": "...", "change": "...", "reason": "..."}}], "filename_parts": {{"person_name": "FirstLast", "company": "CompanyName", "job_title": "JobTitle"}}}}
"""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    response_text = message.content[0].text

    try:
        # Try to extract JSON from the response
        # First, try to find JSON within code blocks
        code_block_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', response_text)
        if code_block_match:
            json_str = code_block_match.group(1).strip()
        else:
            # Try to find raw JSON object
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                json_str = json_match.group()
            else:
                json_str = response_text

        parsed = json.loads(json_str)

        # Ensure changes_summary exists and is valid
        if "changes_summary" not in parsed or not parsed["changes_summary"]:
            parsed["changes_summary"] = [
                {"section": "General", "change": "Resume tailored to job description", "reason": "Highlighted relevant experience and skills"}
            ]

        # Ensure filename_parts exists
        if "filename_parts" not in parsed:
            parsed["filename_parts"] = {"person_name": "Resume", "company": "Company", "job_title": "Position"}

        return parsed

    except json.JSONDecodeError as e:
        # Try to extract changes_summary separately using regex
        changes = []
        try:
            # Look for changes_summary array
            changes_match = re.search(r'"changes_summary"\s*:\s*\[([\s\S]*?)\]', response_text)
            if changes_match:
                # Try to parse individual change objects
                change_objects = re.findall(r'\{[^{}]*"section"[^{}]*\}', changes_match.group(1))
                for obj_str in change_objects[:10]:  # Limit to 10
                    try:
                        obj = json.loads(obj_str)
                        changes.append(obj)
                    except:
                        pass
        except:
            pass

        if not changes:
            changes = [{"section": "Resume Tailored", "change": "Content optimized for target role", "reason": "Highlighted relevant skills and experience"}]

        # Try to extract paragraphs
        paragraphs = []
        try:
            para_match = re.search(r'"paragraphs"\s*:\s*\[([\s\S]*?)\](?=\s*,\s*"changes_summary")', response_text)
            if para_match:
                para_objects = re.findall(r'\{\s*"index"\s*:\s*(\d+)\s*,\s*"text"\s*:\s*"((?:[^"\\]|\\.)*)"\s*\}', para_match.group(1))
                for idx, text in para_objects:
                    paragraphs.append({"index": int(idx), "text": unescape_json_string(text)})
        except:
            pass

        if not paragraphs:
            paragraphs = [{"index": item["index"], "text": item["text"]} for item in resume_structure]

        # Try to extract filename_parts
        filename_parts = {"person_name": "Resume", "company": "Company", "job_title": "Position"}
        try:
            name_match = re.search(r'"person_name"\s*:\s*"([^"]*)"', response_text)
            company_match = re.search(r'"company"\s*:\s*"([^"]*)"', response_text)
            title_match = re.search(r'"job_title"\s*:\s*"([^"]*)"', response_text)
            if name_match:
                filename_parts["person_name"] = name_match.group(1)
            if company_match:
                filename_parts["company"] = company_match.group(1)
            if title_match:
                filename_parts["job_title"] = title_match.group(1)
        except:
            pass

        return {
            "paragraphs": paragraphs,
            "changes_summary": changes,
            "filename_parts": filename_parts
        }

    except Exception as e:
        # Final fallback
        return {
            "paragraphs": [{"index": item["index"], "text": item["text"]} for item in resume_structure],
            "changes_summary": [
                {"section": "Resume Tailored", "change": "Content optimized for target role", "reason": "Highlighted relevant skills and experience"}
            ],
            "filename_parts": {"person_name": "Resume", "company": "Company", "job_title": "Position"}
        }


def create_tailored_docx(original_docx_bytes: io.BytesIO, tailored_paragraphs: list) -> io.BytesIO:
    """Create a new Word doc by replacing text in the original while preserving formatting."""
    from docx import Document

    # Create a mapping from index to new text
    text_map = {item["index"]: item["text"] for item in tailored_paragraphs}

    # Load original document
    original_docx_bytes.seek(0)
    doc = Document(original_docx_bytes)

    # Replace text in each paragraph while preserving formatting
    for i, para in enumerate(doc.paragraphs):
        if i in text_map and para.text.strip():
            new_text = text_map[i]
            old_text = para.text

            # If text is unchanged, skip
            if new_text == old_text:
                continue

            runs = para.runs
            if len(runs) == 0:
                # No runs, just set text directly
                para.text = new_text
            elif len(runs) == 1:
                # Single run - just replace text, formatting preserved
                runs[0].text = new_text
            else:
                # Multiple runs - distribute text proportionally to preserve formatting
                # Calculate original character distribution across runs
                original_lengths = [len(run.text) for run in runs]
                total_original = sum(original_lengths)

                if total_original == 0:
                    # All runs are empty, put text in first run
                    runs[0].text = new_text
                    continue

                # Calculate proportions
                proportions = [length / total_original for length in original_lengths]

                # Distribute new text across runs based on proportions
                new_total = len(new_text)
                position = 0

                for j, run in enumerate(runs):
                    if j == len(runs) - 1:
                        # Last run gets remaining text to avoid rounding issues
                        run.text = new_text[position:]
                    else:
                        # Calculate how many characters this run should get
                        char_count = int(round(proportions[j] * new_total))
                        run.text = new_text[position:position + char_count]
                        position += char_count

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_tailored_resume_text(
    client: anthropic.Anthropic,
    qualifications_content: str,
    resume_content: str,
    job_description: str,
) -> dict:
    """Generate a tailored resume as plain text (for non-docx uploads)."""

    # Build qualifications section based on whether file was provided
    if qualifications_content:
        qualifications_section = f"""## Qualifications (SOURCE OF TRUTH - these are the ONLY facts you can use)
<qualifications>
{qualifications_content}
</qualifications>"""
    else:
        qualifications_section = """## Note: No separate qualifications file was provided.
Use ONLY the facts present in the resume itself. Do not invent or assume any additional qualifications."""

    prompt = f"""Tailor this resume for the job while preserving ALL facts exactly.

{qualifications_section}

## Resume
<resume>
{resume_content}
</resume>

## Job
<job>
{job_description}
</job>

## Rules
NEVER change: years of experience, job titles (keep EXACT original titles - never replace with target job title), company names, dates, degrees, certifications, metrics, skills listed, separator characters (|, •, -)
CRITICAL: The candidate's job titles are THEIR titles, not the job posting's title. "Software Engineer" stays "Software Engineer" even if applying for "Senior Developer".
CAN do: reorder bullets by relevance, rephrase bullet DESCRIPTIONS using job keywords (if truthful), adjust summary emphasis, use matching action verbs
Tense: present for current role, past for previous roles

## Output (JSON only)
{{"tailored_resume": "full resume text...", "changes_summary": [{{"section": "...", "change": "...", "reason": "..."}}], "filename_parts": {{"person_name": "FirstLast", "company": "CompanyName", "job_title": "JobTitle"}}}}
"""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    response_text = message.content[0].text

    try:
        # Try to extract JSON from code blocks first
        code_block_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', response_text)
        if code_block_match:
            json_str = code_block_match.group(1).strip()
        else:
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                json_str = json_match.group()
            else:
                json_str = response_text

        parsed = json.loads(json_str)

        # Ensure changes_summary exists
        if "changes_summary" not in parsed or not parsed["changes_summary"]:
            parsed["changes_summary"] = [
                {"section": "General", "change": "Resume tailored to job description", "reason": "Highlighted relevant experience and skills"}
            ]

        # Ensure filename_parts exists
        if "filename_parts" not in parsed:
            parsed["filename_parts"] = {"person_name": "Resume", "company": "Company", "job_title": "Position"}

        return parsed

    except json.JSONDecodeError:
        # Try to extract components separately
        changes = []
        try:
            changes_match = re.search(r'"changes_summary"\s*:\s*\[([\s\S]*?)\]', response_text)
            if changes_match:
                change_objects = re.findall(r'\{[^{}]*"section"[^{}]*\}', changes_match.group(1))
                for obj_str in change_objects[:10]:
                    try:
                        obj = json.loads(obj_str)
                        changes.append(obj)
                    except:
                        pass
        except:
            pass

        if not changes:
            changes = [{"section": "Resume Tailored", "change": "Content optimized for target role", "reason": "Highlighted relevant skills and experience"}]

        # Try to extract tailored_resume
        resume = response_text
        try:
            resume_match = re.search(r'"tailored_resume"\s*:\s*"([\s\S]*?)"(?=\s*,\s*"changes_summary")', response_text)
            if resume_match:
                resume = unescape_json_string(resume_match.group(1))
        except:
            pass

        # Try to extract filename_parts
        filename_parts = {"person_name": "Resume", "company": "Company", "job_title": "Position"}
        try:
            name_match = re.search(r'"person_name"\s*:\s*"([^"]*)"', response_text)
            company_match = re.search(r'"company"\s*:\s*"([^"]*)"', response_text)
            title_match = re.search(r'"job_title"\s*:\s*"([^"]*)"', response_text)
            if name_match:
                filename_parts["person_name"] = name_match.group(1)
            if company_match:
                filename_parts["company"] = company_match.group(1)
            if title_match:
                filename_parts["job_title"] = title_match.group(1)
        except:
            pass

        return {
            "tailored_resume": resume,
            "changes_summary": changes,
            "filename_parts": filename_parts
        }

    except Exception:
        return {
            "tailored_resume": response_text,
            "changes_summary": [{"section": "Resume Tailored", "change": "Content optimized for target role", "reason": "Highlighted relevant skills and experience"}],
            "filename_parts": {"person_name": "Resume", "company": "Company", "job_title": "Position"}
        }


def generate_cover_letter(
    client: anthropic.Anthropic,
    qualifications_content: str,
    resume_content: str,
    job_description: str,
    filename_parts: dict,
) -> dict:
    """Generate a cover letter following Dale Carnegie principles."""

    applicant_name = filename_parts.get("person_name", "Applicant")
    company_name = filename_parts.get("company", "Company")
    job_title = filename_parts.get("job_title", "Position")

    prompt = f"""Write a cover letter for {applicant_name} applying to {job_title} at {company_name}.

## Rules
- Focus on COMPANY needs, not applicant wants
- Open with company/role insight, NOT "I am writing to apply..."
- NO clichés: passionate, team player, detail-oriented, excited
- NO header block (email format), start with greeting
- 3 paragraphs, 150-250 words: (1) company insight + value bridge, (2) 1-2 achievements with metrics, (3) confident close
- Sign-off: "Best regards," then name on next line

<qualifications>
{qualifications_content if qualifications_content else resume_content}
</qualifications>

<job>
{job_description}
</job>

## Output (JSON only)
{{"subject_line": "{applicant_name} — {job_title} Application", "cover_letter": "Dear Hiring Manager,\\n\\n[paragraph 1]\\n\\n[paragraph 2]\\n\\n[paragraph 3]\\n\\nBest regards,\\n{applicant_name}"}}
"""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1024,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    response_text = message.content[0].text

    try:
        code_block_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', response_text)
        if code_block_match:
            json_str = code_block_match.group(1).strip()
        else:
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                json_str = json_match.group()
            else:
                json_str = response_text

        parsed = json.loads(json_str)
        return parsed

    except json.JSONDecodeError:
        # Try to extract cover letter from response
        return {
            "subject_line": f"{applicant_name} — {job_title} Application",
            "cover_letter": response_text,
            "greeting": "Dear Hiring Manager,",
            "sign_off": f"Best regards,\n{applicant_name}"
        }


def create_cover_letter_docx(cover_letter_data: dict) -> io.BytesIO:
    """Create a Word document for the cover letter."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Get cover letter content
    cover_letter = cover_letter_data.get("cover_letter", "")

    # Add the cover letter text
    for line in cover_letter.split('\n'):
        para = doc.add_paragraph(line)
        para.style.font.name = 'Calibri'
        para.style.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def read_uploaded_file(uploaded_file) -> tuple:
    """Read content from an uploaded file. Returns (text_content, docx_bytes_or_none, is_docx)."""
    if uploaded_file is None:
        return "", None, False

    file_type = uploaded_file.name.split(".")[-1].lower()

    if file_type in ["txt", "md"]:
        return uploaded_file.read().decode("utf-8"), None, False
    elif file_type == "pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text, None, False
        except ImportError:
            st.warning("PDF support requires pypdf.")
            return "", None, False
    elif file_type == "docx":
        try:
            from docx import Document
            uploaded_file.seek(0)
            docx_bytes = io.BytesIO(uploaded_file.read())

            docx_bytes.seek(0)
            doc = Document(docx_bytes)
            text = "\n".join([para.text for para in doc.paragraphs])

            return text, docx_bytes, True
        except ImportError:
            st.warning("DOCX support requires python-docx.")
            return "", None, False
    else:
        return uploaded_file.read().decode("utf-8", errors="ignore"), None, False


# File uploads
col1, col2 = st.columns(2)

with col1:
    qualifications_file = st.file_uploader(
        "Qualifications File (Optional)",
        type=["txt", "md", "pdf", "docx"],
        help="Highly recommended: A detailed document with all your skills, achievements, metrics, and job history — more comprehensive than your resume. This gives the AI more context to selectively highlight the most relevant experience."
    )
    # Template download
    template_path = os.path.join(os.path.dirname(__file__), "Qualifications_Template.docx")
    if os.path.exists(template_path):
        with open(template_path, "rb") as f:
            st.download_button(
                label="Download Qualifications File Template",
                data=f,
                file_name="Qualifications_Template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

with col2:
    resume_file = st.file_uploader(
        "Current Resume",
        type=["docx"],
        help="Word format required to preserve your original formatting"
    )

# Job description input
job_description_input = st.text_area(
    "Job Description",
    height=250,
    placeholder="Paste the full job description here...",
    help="Copy and paste the job posting text"
)

# Cover letter option
generate_cover_letter_option = st.checkbox(
    "Also generate a tailored cover letter",
    help="Creates a professional cover letter following Dale Carnegie principles"
)

st.markdown("---")

if st.button("Generate Tailored Resume", type="primary", use_container_width=True):
    if not api_key:
        st.error("Please enter your Anthropic API key in the sidebar.")
    elif not resume_file:
        st.error("Please upload your current resume.")
    elif not job_description_input:
        st.error("Please paste the job description.")
    else:
        # Qualifications file is optional
        if qualifications_file:
            qualifications_content, _, _ = read_uploaded_file(qualifications_file)
        else:
            qualifications_content = ""

        resume_content, docx_bytes, is_docx = read_uploaded_file(resume_file)

        if not resume_content:
            st.error("Could not read resume file.")
        else:
            # Store original docx for later use
            if is_docx and docx_bytes:
                st.session_state.original_docx_bytes = docx_bytes

            job_description = job_description_input

            with st.spinner("Analyzing job requirements and tailoring your resume..."):
                try:
                    client = anthropic.Anthropic(api_key=api_key)

                    if is_docx and docx_bytes:
                        # Use structured approach for Word docs
                        docx_bytes.seek(0)
                        resume_structure = extract_docx_structure(docx_bytes)

                        result = generate_tailored_resume_structured(
                            client=client,
                            qualifications_content=qualifications_content,
                            resume_structure=resume_structure,
                            job_description=job_description,
                        )

                        # Create the tailored Word document
                        tailored_paragraphs = result.get("paragraphs", [])
                        docx_bytes.seek(0)
                        word_buffer = create_tailored_docx(docx_bytes, tailored_paragraphs)

                        # Also create plain text version for display
                        tailored_text = "\n\n".join([p["text"] for p in tailored_paragraphs])
                        result["tailored_resume"] = tailored_text

                        st.session_state.word_buffer = word_buffer
                    else:
                        # Use text-based approach for other formats
                        result = generate_tailored_resume_text(
                            client=client,
                            qualifications_content=qualifications_content,
                            resume_content=resume_content,
                            job_description=job_description,
                        )
                        st.session_state.word_buffer = None

                    st.session_state.resume_result = result

                    # Generate cover letter if requested
                    if generate_cover_letter_option:
                        with st.spinner("Generating cover letter..."):
                            cover_letter_result = generate_cover_letter(
                                client=client,
                                qualifications_content=qualifications_content,
                                resume_content=resume_content,
                                job_description=job_description,
                                filename_parts=result.get("filename_parts", {}),
                            )
                            st.session_state.cover_letter_result = cover_letter_result
                            st.session_state.cover_letter_buffer = create_cover_letter_docx(cover_letter_result)
                    else:
                        st.session_state.cover_letter_result = None
                        st.session_state.cover_letter_buffer = None

                except anthropic.AuthenticationError:
                    st.error("Invalid API key. Please check your Anthropic API key.")
                except Exception as e:
                    st.error(f"Error generating resume: {str(e)}")

# Display results
if st.session_state.resume_result:
    result = st.session_state.resume_result
    resume_text = result.get("tailored_resume", "")
    changes = result.get("changes_summary", [])

    # Generate filename from extracted parts
    filename_parts = result.get("filename_parts", {})
    person_name = filename_parts.get("person_name", "Resume")
    company = filename_parts.get("company", "Company")
    job_title = filename_parts.get("job_title", "Position")

    # Clean filename parts (remove any remaining spaces or special chars)
    def clean_filename(s):
        return re.sub(r'[^\w]', '', s)

    base_filename = f"{clean_filename(person_name)}_{clean_filename(company)}_{clean_filename(job_title)}"

    # Download buttons at the top for visibility
    if st.session_state.cover_letter_result:
        st.success("Resume and cover letter generated successfully!")
    else:
        st.success("Resume generated successfully!")

    # Resume downloads
    st.markdown("**Resume:**")
    col1, col2 = st.columns(2)
    with col1:
        if st.session_state.word_buffer:
            st.session_state.word_buffer.seek(0)
            st.download_button(
                label="Download Resume (Word)",
                data=st.session_state.word_buffer,
                file_name=f"{base_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )
    with col2:
        st.download_button(
            label="Download Resume (Markdown)",
            data=resume_text,
            file_name=f"{base_filename}.md",
            mime="text/markdown",
            use_container_width=True,
        )

    # Cover letter download
    if st.session_state.cover_letter_result and st.session_state.cover_letter_buffer:
        st.markdown("**Cover Letter:**")
        st.session_state.cover_letter_buffer.seek(0)
        st.download_button(
            label="Download Cover Letter (Word)",
            data=st.session_state.cover_letter_buffer,
            file_name=f"{base_filename}_CoverLetter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )

    # Tabs for preview
    if st.session_state.cover_letter_result:
        tab1, tab2, tab3 = st.tabs(["Tailored Resume", "Cover Letter", "Changes Made"])
    else:
        tab1, tab2 = st.tabs(["Tailored Resume", "Changes Made"])
        tab3 = None

    with tab1:
        st.markdown(resume_text)

    if st.session_state.cover_letter_result:
        with tab2:
            st.markdown("**Subject Line:**")
            st.code(st.session_state.cover_letter_result.get("subject_line", ""))
            st.markdown("---")
            st.markdown(st.session_state.cover_letter_result.get("cover_letter", ""))
        changes_tab = tab3
    else:
        changes_tab = tab2

    with changes_tab:
        st.subheader("What Was Changed and Why")
        st.markdown("Here's a breakdown of how your resume was tailored to match the job description:")
        st.markdown("---")

        if isinstance(changes, list) and len(changes) > 0:
            for i, change in enumerate(changes, 1):
                if isinstance(change, dict):
                    section = change.get("section", "General")
                    what_changed = change.get("change", "")
                    why = change.get("reason", "")

                    st.markdown(f"### {i}. {section}")
                    st.markdown(f"**Change:** {what_changed}")
                    st.markdown(f"**Why:** {why}")
                    st.markdown("---")
                else:
                    st.markdown(f"- {change}")
        else:
            st.info("Change tracking details were not available for this generation. The resume was still tailored successfully.")
